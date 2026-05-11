from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
import os
from datetime import datetime


# -------------------------------------------------
# Normalize Markdown (**Title**, **Section**) → #, ##
# -------------------------------------------------
def normalize_markdown(text: str) -> str:
    lines = text.split("\n")
    normalized = []

    for i, line in enumerate(lines):
        stripped = line.strip()

        # MAIN TITLE
        if i == 0 and stripped.startswith("**") and stripped.endswith("**"):
            normalized.append("# " + stripped.replace("**", ""))
            continue

        # ALL SECTION HEADERS
        if stripped.startswith("**") and stripped.endswith("**"):
            normalized.append("## " + stripped.replace("**", ""))
            continue

        normalized.append(line)

    return "\n".join(normalized)


# -------------------------------------------------
# DOCX GENERATOR (PRO STYLE – IMAGE 1)
# -------------------------------------------------
def generate_docx(jd_text: str) -> BytesIO:
    jd_text = normalize_markdown(jd_text)

    buffer = BytesIO()
    doc = Document()

    for line in jd_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # MAIN TITLE
        if line.startswith("# "):
            h = doc.add_heading(line.replace("# ", ""), level=0)
            run = h.runs[0]
            run.font.size = Pt(26)  # Increased from 22
            run.bold = True
            h.paragraph_format.space_after = Pt(16)

        # SECTION HEADERS
        elif line.startswith("## "):
            h = doc.add_heading(line.replace("## ", ""), level=1)
            run = h.runs[0]
            run.font.size = Pt(16)  # Increased from 14
            run.bold = True
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(8)

        # BULLETS
        elif line.startswith(("•", "-", "*")):
            p = doc.add_paragraph(line.replace("•", "").strip(), style="List Bullet")
            p.paragraph_format.space_after = Pt(4)

        # NORMAL TEXT
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -------------------------------------------------
# PDF GENERATOR
# -------------------------------------------------
def generate_pdf(jd_text: str) -> BytesIO:
    jd_text = normalize_markdown(jd_text)
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    title = ParagraphStyle(
        "Title",
        fontName="Helvetica-Bold",
        fontSize=26,  # Increased from 20
        spaceAfter=16,  # Increased from 14
        textColor="navy"  # Optional: match Image 2 color
    )

    section = ParagraphStyle(
        "Section",
        fontName="Helvetica-Bold",
        fontSize=16,  # Increased from 14
        spaceBefore=12,  # Increased from 10
        spaceAfter=8  # Increased from 6
    )

    body = ParagraphStyle(
        "Body",
        fontSize=11,
        leading=14,
        spaceAfter=6,
        alignment=TA_LEFT
    )

    bullet = ParagraphStyle(
        "Bullet",
        fontSize=11,
        leftIndent=14,
        spaceAfter=4
    )

    elements = []

    for line in jd_text.split("\n"):
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 6))
            continue

        if line.startswith("# "):
            elements.append(Paragraph(line.replace("# ", ""), title))
        elif line.startswith("## "):
            elements.append(Paragraph(line.replace("## ", ""), section))
        elif line.startswith(("•", "-", "*")):
            elements.append(Paragraph(line.replace("•", ""), bullet))
        else:
            elements.append(Paragraph(line, body))

    doc.build(elements)
    buffer.seek(0)
    return buffer

# -------------------------------------------------
# EXPORT FUNCTIONS (USED BY STREAMLIT)
# -------------------------------------------------
def export_to_docx(jd_text: str, filename: str) -> str:
    buffer = generate_docx(jd_text)

    output_dir = "exports"
    os.makedirs(output_dir, exist_ok=True)

    path = os.path.join(
        output_dir,
        f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    with open(path, "wb") as f:
        f.write(buffer.getvalue())

    return path


def export_to_pdf(jd_text: str, filename: str) -> str:
    buffer = generate_pdf(jd_text)

    output_dir = "exports"
    os.makedirs(output_dir, exist_ok=True)

    path = os.path.join(
        output_dir,
        f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )

    with open(path, "wb") as f:
        f.write(buffer.getvalue())

    return path
