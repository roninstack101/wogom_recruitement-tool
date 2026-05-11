# app/api/jd.py

from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from app.agents.jd_generator import generate_jd
from app.agents.jd_clarifier import generate_clarifying_questions
from app.agents.jd_chatbot import refine_jd
from app.agents.profile_builder import build_profile
from app.agents.role_suggester import suggest_roles
from app.db.database import get_db
from app.db.models import JDFormData
import json

router = APIRouter(
    prefix="/jd",
    tags=["Job Description"]
)


# ── Helper ─────────────────────────────────────────────

def _form_row_to_dict(row: JDFormData) -> dict:
    return {
        "id": row.id,
        "role": row.role,
        "department": row.department,
        "location": row.location or "",
        "employment_type": row.employment_type or "Full-time",
        "work_mode": row.work_mode or "",
        "travel_required": row.travel_required or "",
        "reporting_to": row.reporting_to or "",
        "experience": row.experience or "",
        "minimum_education": row.minimum_education or "",
        "salary": row.salary or "",
        "urgency": row.urgency or "",
        "new_or_scaling": row.new_or_scaling or "",
        "must_have_skills": row.must_have_skills or "",
        "other_skills": row.other_skills or "",
        "key_responsibilities": row.key_responsibilities or "",
        "generated_jd": row.generated_jd or "",
        "generated_profile": row.generated_profile or "",
        "created_at": row.created_at.isoformat() if row.created_at else None,
    }


# ── Saved Forms (replaces Google Sheets) ───────────────

@router.get("/forms")
def list_saved_forms(db: Session = Depends(get_db)):
    """List all previously saved JD intake forms from the database."""
    rows = db.query(JDFormData).order_by(JDFormData.created_at.desc()).all()
    return [_form_row_to_dict(r) for r in rows]


@router.post("/forms")
def save_form(payload: dict, db: Session = Depends(get_db)):
    """Save a new JD intake form to the database."""
    form = JDFormData(
        role=payload.get("role", "").strip(),
        department=payload.get("department", "").strip(),
        location=payload.get("location", "").strip(),
        employment_type=payload.get("employment_type", "Full-time").strip(),
        work_mode=payload.get("work_mode", "").strip(),
        travel_required=payload.get("travel_required", "").strip(),
        reporting_to=payload.get("reporting_to", "").strip(),
        experience=payload.get("experience", "").strip(),
        minimum_education=payload.get("minimum_education", "").strip(),
        salary=payload.get("salary", "").strip(),
        urgency=payload.get("urgency", "").strip(),
        new_or_scaling=payload.get("new_or_scaling", "").strip(),
        must_have_skills=payload.get("must_have_skills", "").strip(),
        other_skills=payload.get("other_skills", "").strip(),
        key_responsibilities=payload.get("key_responsibilities", "").strip(),
    )
    db.add(form)
    db.commit()
    db.refresh(form)
    return _form_row_to_dict(form)


@router.put("/forms/{form_id}/jd")
def update_form_jd(form_id: int, payload: dict, db: Session = Depends(get_db)):
    """Update the generated JD text on a saved form."""
    form = db.query(JDFormData).filter(JDFormData.id == form_id).first()
    if not form:
        return {"error": "Form not found"}
    form.generated_jd = payload.get("generated_jd", "")
    db.commit()
    db.refresh(form)
    return _form_row_to_dict(form)


@router.put("/forms/{form_id}/profile")
def update_form_profile(form_id: int, payload: dict, db: Session = Depends(get_db)):
    """Save the AI-generated candidate profile on a saved form."""
    form = db.query(JDFormData).filter(JDFormData.id == form_id).first()
    if not form:
        return {"error": "Form not found"}
    import json
    profile_data = payload.get("generated_profile")
    form.generated_profile = json.dumps(profile_data) if isinstance(profile_data, dict) else (profile_data or "")
    db.commit()
    db.refresh(form)
    return _form_row_to_dict(form)


@router.delete("/forms/{form_id}")
def delete_form(form_id: int, db: Session = Depends(get_db)):
    """Delete a saved JD intake form."""
    form = db.query(JDFormData).filter(JDFormData.id == form_id).first()
    if not form:
        return {"error": "Form not found"}
    db.delete(form)
    db.commit()
    return {"ok": True}


# ── Legacy /roles endpoint (now reads from DB) ────────

@router.get("/roles")
def get_roles(db: Session = Depends(get_db)):
    """Return saved forms as roles (backward compat)."""
    rows = db.query(JDFormData).order_by(JDFormData.created_at.desc()).all()
    return [_form_row_to_dict(r) for r in rows]

@router.post("/clarify")
def clarify_jd_api(payload: dict):
    """Agent 1: Generate clarifying questions from form data (no draft needed)."""
    questions = generate_clarifying_questions(form_data=payload)

    if isinstance(questions, str):
        questions = json.loads(questions)

    return {"questions": questions}


@router.post("/profile")
def profile_builder_api(payload: dict):
    """Agent 2: Build ideal candidate profile."""
    profile = build_profile(
        form_data=payload.get("form_data", {}),
        clarification_answers=payload.get("answers", [])
    )
    return {"profile": profile}


@router.post("/suggest-roles")
def suggest_roles_api(payload: dict):
    """Suggest alternative role names from the profile."""
    profile = payload.get("profile", {})
    instruction = payload.get("instruction", None)
    suggestions = suggest_roles(profile, instruction)
    return {"suggestions": suggestions}


@router.post("/generate")
def generate_jd_api(payload: dict):
    """Agent 3: Generate JD from profile + form data."""
    jd = generate_jd(
        form_data=payload.get("form_data", payload),
        profile=payload.get("profile", None)
    )
    return {"jd": jd}


@router.post("/refine")
def refine_jd_api(payload: dict):
    """Agent 4: Refine JD based on user instruction."""
    updated_jd = refine_jd(
        current_jd=payload["jd"],
        instruction=payload["instruction"],
        role=payload.get("role", ""),
        session_id=payload.get("session_id", "")
    )
    return {"jd": updated_jd}


@router.post("/export-docx")
def export_jd_docx(payload: dict):
    """Export JD as a formatted .docx file."""
    import io
    import re
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    jd_text = payload.get("jd", "")
    role = payload.get("role", "Job Description")

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── Parse markdown lines ──
    lines = jd_text.split("\n")

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # Heading 1: # Title
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading = stripped[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            continue

        # Heading 2: ## Section
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            p = doc.add_heading(heading, level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x5e)
                run.font.size = Pt(13)
            continue

        # Heading 3: ### Sub-section
        if stripped.startswith("### "):
            heading = stripped[4:].strip()
            p = doc.add_heading(heading, level=3)
            for run in p.runs:
                run.font.size = Pt(12)
            continue

        # Bullet points: - item or * item
        if re.match(r"^[-*•]\s+", stripped):
            text = re.sub(r"^[-*•]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold markers **text**
            _add_formatted_run(p, text)
            continue

        # Numbered list: 1. item
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, text)
            continue

        # Horizontal rules
        if re.match(r"^[-─═]{3,}$", stripped):
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)

    # ── Save to buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{role.replace(' ', '_')}_JD.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


def _add_formatted_run(paragraph, text: str):
    """Add text to a paragraph, converting **bold** markers to bold runs."""
    import re
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

